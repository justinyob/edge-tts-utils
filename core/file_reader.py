import logging
import os

import chardet
import docx

from core.exceptions import FileReadError

__all__ = ["FileReadError", "read_file"]

log = logging.getLogger(__name__)


def _read_txt(path: str) -> str:
    try:
        with open(path, "rb") as f:
            sample = f.read(10_000)
        detected = chardet.detect(sample) if sample else {"encoding": None}
        encoding = detected.get("encoding") or "utf-8"
        try:
            with open(path, "r", encoding=encoding, errors="strict") as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
    except OSError as e:
        raise FileReadError(f"Could not open text file: {e}") from e


def _read_docx(path: str) -> str:
    try:
        document = docx.Document(path)
    except Exception as e:
        raise FileReadError(f"Could not open Word document: {e}") from e

    table_paragraph_ids: set[int] = set()
    try:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        table_paragraph_ids.add(id(p))
    except Exception as e:
        raise FileReadError(f"Failed reading document tables: {e}") from e

    paragraphs: list[str] = []
    try:
        for p in document.paragraphs:
            if id(p) in table_paragraph_ids:
                continue
            text = (p.text or "").strip()
            if text:
                paragraphs.append(text)
    except Exception as e:
        raise FileReadError(f"Failed reading document paragraphs: {e}") from e

    return "\n\n".join(paragraphs)


def read_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return _read_txt(path)
    if ext == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file type: {ext}. Supported: .txt, .docx")


if __name__ == "__main__":
    import tempfile

    # Test 1: .txt round-trip
    content = "Hello, world.\nThis is a test file.\nLine three."
    fd, txt_path = tempfile.mkstemp(suffix=".txt")
    os.close(fd)
    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(content)
        result = read_file(txt_path)
        assert result == content, f"txt mismatch: {result!r} vs {content!r}"
        print("  test 1 (.txt): OK")
    finally:
        os.remove(txt_path)

    # Test 2: .docx with 3 paragraphs
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        d = docx.Document()
        d.add_paragraph("Paragraph one — the beginning.")
        d.add_paragraph("Paragraph two — the middle.")
        d.add_paragraph("")  # empty, should be skipped
        d.add_paragraph("Paragraph three — the end.")
        d.save(docx_path)

        result = read_file(docx_path)
        assert "Paragraph one" in result
        assert "Paragraph two" in result
        assert "Paragraph three" in result
        # Empty paragraph should not produce triple newline
        assert "\n\n\n" not in result
        print("  test 2 (.docx): OK")
    finally:
        os.remove(docx_path)

    # Test 3: unsupported extension
    try:
        read_file("/tmp/whatever.pdf")
    except ValueError as e:
        assert ".pdf" in str(e)
        print("  test 3 (.pdf rejected): OK")
    else:
        raise AssertionError("expected ValueError for .pdf")

    print("FileReader: OK")
